import streamlit as st
import streamlit.components.v1 as components
from openai import OpenAI
import json
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches
import plotly.graph_objects as go
import plotly.express as px
import plotly.io as pio
import pandas as pd
import os
import urllib.parse

# Fonction pour créer un document Word
def make_docx(title: str, content: str) -> bytes:
    buf = io.BytesIO()
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M"))
    for line in content.splitlines():
        doc.add_paragraph(line)
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# Fonction pour exporter les scores en CSV
def make_scores_csv(scores: dict) -> str:
    lines = ["competence,score"]
    for comp, score in scores.items():
        lines.append(f"{comp},{score:.2f}")
    return "\n".join(lines)

# Directive de langue pour les réponses (Français / Wolof)
def get_lang_directive() -> str:
    lang = st.session_state.get('app_lang', 'Français')
    if lang == 'Wolof':
        # Directive simple et robuste pour forcer la langue Wolof
        return "Réponds uniquement en Wolof (langue wolof standard du Sénégal)."
    return "Réponds en français."

# Mini-système de traduction pour l'UI (Français / Wolof)
TRANSLATIONS = {
    'Français': {
        'tab_eval': "Évaluation",
        'tab_results': "Résultats",
        'tab_reco': "Recommandations",
        'tab_adja': "Coach Fatouma",
          'app_title': "🇸🇳 Outil de Profilage entrepreneuriale",
          'app_tagline': "Évaluez vos compétences entrepreneuriales et obtenez des recommandations personnalisées",
          'adja_caption': "Fatouma répond aux questions sur l'entrepreneuriat.",
        'tab1_header': "Évaluation des Compétences",
        'tab1_rubriques': "Rubriques d'évaluation",
        'tab1_instruction': "Sélectionnez une rubrique puis évaluez chaque affirmation sur une échelle de 1 (Pas du tout d'accord) à 5 (Tout à fait d'accord)",
        'progress_global': "Progression globale",
        'sidebar_info': "📋 Informations",
        'sidebar_name': "Nom complet",
        'sidebar_age': "Âge",
        'sidebar_sector': "Secteur d'activité",
        'sidebar_sector_custom': "Secteur personnalisé",
        'sidebar_sector_placeholder': "Saisissez votre secteur d'activité",
        'sidebar_experience': "Expérience entrepreneuriale",
        'sidebar_language': "Langue",
        'questions_answered': "Questions répondues: {answered}/{total}",
        'profile_calculated': "✅ Votre Profil a été Calculé",
        'complete_info': "Complétez vos informations pour afficher votre profil d’entrepreneur",
        'next_rubrique_button': "Rubrique suivante",
        'resume_rapide': "🔎 Résumé rapide",
        'score_global': "📈 Score Global",
        'points_forts': "⭐ Points Forts",
        'axes_amelioration': "⚠️ Axes d'Amélioration",
        'delta_excellent': "Excellent!",
        'delta_to_develop': "À développer",
        'open_results_hint': "Pour le détail complet, ouvrez l'onglet \"📊 Résultats\".",
        'monter': "⬆️ Monter",
        'nav_success_heading': "🎉 Félicitations ! Votre évaluation est terminée.",
        'nav_results_hint': "👉 Consultez maintenant l'onglet \"📊 Résultats\" pour voir votre profil détaillé",
        'nav_reco_hint': "💡 Puis l'onglet \"💡 Recommandations\" pour obtenir des conseils personnalisés",
        'nav_all_completed': "✅ Toutes les compétences sont évaluées !",
        'nav_complete_personal_info': "📝 Complétez vos informations personnelles ci-dessus pour générer votre profil",
        'nav_progress_label': "📊 Progression : {percent}%",
        'nav_continue_eval': "🎯 Continuez à évaluer les compétences pour débloquer vos résultats",
        'company_name_label': "Nom de l'entreprise",
        'non_renseigne': "Non renseigné",
        'results_header': "📊 Votre Profil Entrepreneurial",
        'download_txt': "💾 Télécharger en TXT",
        'download_word': "📄 Télécharger en Word",
        'generating': "Génération en cours...",
        'mentorat_button': "👥 Recommandations de Mentorat",
        'financement_button': "💼 Opportunités de Financement",
        'plan_action_90_title': "🗓️ Plan d'action 90 jours",
        'plan_action_90_generate': "🗓️ Générer le plan 90 jours",
        'analyse_complete_button': "🚀 Analyse Complète et Recommandations Globales",
        'download_analysis_complete': "💾 Télécharger l'analyse complète",
        'download_analysis_word': "Télécharger en Word (.docx)",
        'no_resource_match': "Aucune ressource correspondante. Essayez un autre mot-clé.",
        'journal_coaching_title': "📝 Journal de Coaching",
        'download_journal_csv': "💾 Télécharger Journal (CSV)",
        'journal_empty_caption': "Le journal de coaching est vide pour le moment.",
        'adja_profile_success': "✅ Ton profil est pris en compte par Fatouma pour des conseils personnalisés.",
          'adja_info_prompt': "ℹ️ Pour des conseils plus personnalisés, complète l’onglet ‘Évaluation’.",
        'goto_eval_button': "Aller à l’onglet Évaluation",
        'goto_eval_warning': "Clique sur l’onglet ‘Évaluation’ en haut de la page pour commencer.",
        'radar_trace_name': "Vos Compétences",
        'score_label': "Score",
        'footer_tool_heading': "🌍 Outil de Profilage Entrepreneurial - Sénégal",
        'footer_tool_sub': "Développé par M-T pour accompagner les entrepreneurs sénégalais",
        'footer_credit_by': "@Développé par Moctar TALL",
        'footer_rights': "All Rights Reserved",
        'footer_phone_label': "📞 Tél :",
        'to_evaluate': "À évaluer",
        'actions_recommandees': "🎯 Actions Recommandées",
        'vous_etes_ici': "VOUS ÊTES ICI",
        'local_resources_title': "📚 Ressources Locales",
        'search_resources_placeholder': "Rechercher une ressource (ex: financement, formation, mentorat)",
        'share_whatsapp': "Partager via WhatsApp",
        'doc_title_financement': "Opportunités de Financement",
        'doc_title_mentorat': "Recommandations de Mentorat",
        'doc_title_analyse_complete': "Analyse complète & Recommandations",
        'click_rubrique_hint': "👆 Cliquez sur une rubrique ci-dessus pour commencer l'évaluation",
        'radar_map_title': "🕸️ Cartographie de vos compétences",
        'heatmap_comp_title': "🔥 Heatmap des compétences",
    },
    'Wolof': {
        'tab_eval': "Seetu Mën-mën yi",
        'tab_results': "njureef",
        'tab_reco': "Ndigël",
        'tab_adja': "Cooc Fatouma",
        'app_title': "Jumtukaay bu seet profilu ëmbëru Senegaal",
        'app_tagline': "Seet sa mën-mën ci entrepreneuriat te am ndigël yu ci sa bopp",
        'adja_caption': "Fatouma dees na tontu laaj yi ci entrepreneuriat rekk.",
        'tab1_header': "Seetu Mën-mën yi",
        'tab1_rubriques': "Lislaasu seetu",
        'tab1_instruction': "Fal benn lislaas, te jéggal benn wax ci tegleel 1 di 5 (1: duñoo noppi, 5: noppi nopp)",
        'progress_global': "Yéene jëm ci yenn ñaari xaal yi",
        'sidebar_info': "📋 Say Xibaar",
        'sidebar_name': "Sa Tur",
        'sidebar_age': "Say At",
        'sidebar_sector': "Sa Sektoru liggéey",
        'sidebar_sector_custom': "Sektor bu sa bopp",
        'sidebar_sector_placeholder': "Bind sektor bu sa liggéey",
        'sidebar_experience': "Xéy ci entrepreneuriat",
        'sidebar_language': "Kalama",
        'questions_answered': "Laaj yi jëggalee: {answered}/{total}",
        'profile_calculated': "✅ Sa profil bi ñu kalkule na",
        'complete_info': "Tammal say xibaar ngir wone sa profil ëmbëru",
        'next_rubrique_button': "Rubrik bu ci topp",
        'resume_rapide': "🔎 Wone bu gaaw",
        'score_global': "📈 Score Biir",
        'points_forts': "⭐ Mën-mën yu am",
        'axes_amelioration': "⚠️ Yoonu soppali",
        'delta_excellent': "Baax na lool!",
        'delta_to_develop': "Wara yokk",
        'open_results_hint': "Ngir gëstu bu mat, ubbil \"📊 njureef\".",
        'monter': "⬆️ Yéeg",
        'nav_success_heading': "🎉 Jàmm rekk! Sa seetu jeex na.",
        'nav_results_hint': "👉 Jëll ci \"📊 njureef\" ngir gis sa profil bu mat",
        'nav_reco_hint': "💡 Ci topp, \"💡 Ndigël\" ngir am ndigël yu ci sa bopp",
        'nav_all_completed': "✅ Mën-mën yépp ñu seet na!",
        'nav_complete_personal_info': "📝 Tammal say xibaar ci kaw ngir génn sa profil",
        'nav_progress_label': "📊 Yéene : {percent}%",
        'nav_continue_eval': "🎯 Kontineel seet mën-mën yi ngir ubbi say njureef",
        'company_name_label': "Turu ëntërpris bi",
        'non_renseigne': "Duñu ko joxe",
        'results_header': "📊 Sa Profil ëmbëru",
        'download_txt': "💾 Yebal ci TXT",
        'download_word': "📄 Yebal ci Word",
        'generating': "Gënn ci def...",
        'mentorat_button': "👥 Ndigël ci Mentoraat",
        'financement_button': "💼 Jariñu Laccas",
        'plan_action_90_title': "🗓️ Palaan 90 fan",
        'plan_action_90_generate': "🗓️ Sos palaan 90 fan",
        'analyse_complete_button': "🚀 Analys bu mat ak Ndigël yu bari",
        'download_analysis_complete': "💾 Yebal analays bi",
        'download_analysis_word': "Yebal ci Word (.docx)",
        'no_resource_match': "Amul resurs bu japp. Jéem benn baat bu wuute.",
        'journal_coaching_title': "📝 Jurnal bu coaching",
        'download_journal_csv': "💾 Yebal Jurnal (CSV)",
        'journal_empty_caption': "Jurnal bu coaching bi des na.",
        'adja_profile_success': "✅ Fatouma dafa jëfandikoo sa profil ngir ndigël yu ci sa bopp.",
          'adja_info_prompt': "ℹ️ Ngir am ndigël yu gën a tekki, seetal onglet ‘Seetu’.",
        'goto_eval_button': "Dellu ci onglet ‘Seetu’",
        'goto_eval_warning': "Seetu onglet ‘Seetu’ ci kaw bi ngir tàmbalee.",
        'radar_trace_name': "Sa Mën‑mën yi",
        'score_label': "Njaaxum",
        'footer_tool_heading': "🌍 Jumtukaay seetu ëmbëru - Senegaal",
        'footer_tool_sub': "Defu ko M‑T ngir tàllal ëmbëru Senegaal",
        'footer_credit_by': "@Moctar TALL moo ko def",
        'footer_rights': "Droit yëpp mooy moom",
        'footer_phone_label': "📞 Téléfoon :",
        'to_evaluate': "Ñaata laaj nga koy jéggal",
        'actions_recommandees': "🎯 Jëf yi ñu jox ndigël",
        'vous_etes_ici': "FOO NEKK",
        'local_resources_title': "📚 Resurs yu dëkk",
        'search_resources_placeholder': "Seet benn resurs (misaal: laccas, jàng, mentoraat)",
        'share_whatsapp': "Sédd ci WhatsApp",
        'doc_title_financement': "Jariñu Laccas",
        'doc_title_mentorat': "Ndigël ci Mentoraat",
        'doc_title_analyse_complete': "Analys bu mat ak Ndigël",
        'click_rubrique_hint': "👆 Bësal benn rubrik ci kaw ngir tàmbalee seetu",
        'radar_map_title': "🕸️ Kaarti sa mën‑mën yi",
        'heatmap_comp_title': "🔥 Màppu‑xeetu mën‑mën yi",
    }
}

# Libellés Wolof pour les compétences (affichage)
COMP_LABELS = {
    'Français': {
        "Leadership": "Leadership",
        "Gestion & Délégation": "Gestion & Délégation",
        "Créativité & Innovation": "Créativité & Innovation",
        "Réseautage & Relations": "Réseautage & Relations",
        "Résilience & Persévérance": "Résilience & Persévérance",
        "Gestion Financière": "Gestion Financière",
    },
    'Wolof': {
        "Leadership": "Jiitu",
        "Gestion & Délégation": "Toppatoo & Jox Njël",
        "Créativité & Innovation": "Yëngu‑yëng & Yeesal",
        "Réseautage & Relations": "Jokkoo & Jàppante",
        "Résilience & Persévérance": "Muñ & Tëxë",
        "Gestion Financière": "Toppatoo Laccas",
    }
}

# Traductions Wolof des questions par rubrique (ordre synchronisé avec COMPETENCES)
COMP_QUESTIONS_TRANSLATIONS = {
    'Wolof': {
        "Leadership": [
            "Damay tàmbali lu ëpp ci njëw gi",
            "Xam naa ni laa taxawale ñépp te may leen xéy",
            "Damay wone sama seen te woyof te doxlu",
            "Xam naa naari dëgg yu metti te def na ko",
            "Damay jox liggéeykat ñu njël te may leen bokk sañ-sañ",
            "Damay wër jàmm te jàppantoo fax ci yéngu‑yëngu"
        ],
        "Gestion & Délégation": [
            "Damay jox sañ-sañ liggéey yi ci sama équipe bu yomb",
            "Damay gëm ñeneen ngir jëfandikoo liggéey yu am solo",
            "Xam naa toppatoo ak plaani bu wér",
            "Mën naa sàmm ay poroje yu bari ci benn jam",
            "Damay setal lu jiitu ak jamono yu jeex",
            "Damay teggle yoon yu topp ngir sàmm jeeg ak yéene"
        ],
        "Créativité & Innovation": [
            "Damay génné xalaat yu bees bu yomb",
            "Begg naa seet yoon yu bees te jéem",
            "Damay tere xaalis bu nekkoon te lajj status quo",
            "Mën naa xool yoon yu am jariñu",
            "Damay suññ xalaat te jëfandikoo ko ci jëf yi",
            "Damay xool marse bi te jàpp ci gaaw ngir soppi xalaat yi"
        ],
        "Réseautage & Relations": [
            "Damay def jàppante yu liggéey bu yomb",
            "Damay sàmm benn réseau bu di dox ci jamono",
            "Damay jëfandikoo sama réseau ngir yenn xél yi",
            "Damay bokk ci mbootal yi ak waa mbir yu bari",
            "Damay sàmm jàppante bu yor yàgg",
            "Damay def ay parteneer yi bare jariñu ci ñaari bañ"
        ],
        "Résilience & Persévérance": [
            "Damay tekki te muñ ci gàddaay yi",
            "Damay dëgër ci yéene yu yàgg",
            "Soo toppoo ma dafañu tax ma dellu gaaw gannaaw li moye",
            "Damay wër lu baax ci xaalis yu metti",
            "Damay sàmm naqar te wér-góor ci dëgg‑dëgg",
            "Damay soppi palaan bi ci jamono bu xëtul te tëgg xéy yi"
        ],
        "Gestion Financière": [
            "Xam naa xew-xew yi ci wàll laccas bu yomb",
            "Mën naa topp sa ñoom laccas ak bidget bu wér",
            "Mën naa seet yoon ngir am laccas",
            "Damay taxawal ci wàll laccas ak xam-xam",
            "Mën naa plaani cash‑flow ci digg‑bopp",
            "Mën naa teg leppi ci naqar bu jariñu te kenn di ko jënd"
        ]
    }
}

def tr_question(comp_name: str, index: int, default: str) -> str:
    """Retourne la question localisée selon la rubrique et l'index."""
    lang = st.session_state.get('app_lang', 'Français')
    if lang == 'Wolof':
        try:
            return COMP_QUESTIONS_TRANSLATIONS['Wolof'][comp_name][index]
        except Exception:
            return default
    return default

def tr_comp(comp_name: str) -> str:
    lang = st.session_state.get('app_lang', 'Français')
    return COMP_LABELS.get(lang, COMP_LABELS['Français']).get(comp_name, comp_name)

def tr(key: str) -> str:
    """Retourne la traduction selon la langue choisie, avec fallback FR."""
    lang = st.session_state.get('app_lang', 'Français')
    return TRANSLATIONS.get(lang, TRANSLATIONS['Français']).get(key, TRANSLATIONS['Français'].get(key, key))

# Définir la langue par défaut sur Français si non choisie
if 'app_lang' not in st.session_state:
    st.session_state['app_lang'] = 'Français'

# Mini référentiel de ressources locales (Sénégal)
LOCAL_RESOURCES = [
    {
        "name": "DER/FJ",
        "tags": ["financement", "accompagnement", "incubation", "jeunes", "femmes"],
        "description": "Délégation générale à l’Entrepreneuriat Rapide des Femmes et des Jeunes — financement, incubation, appui aux jeunes et femmes.",
        "link": "https://der.sn"
    },
    {
        "name": "APIX",
        "tags": ["investissement", "formalisation", "guichet unique"],
        "description": "Promotion des investissements et guichet unique pour création d’entreprise.",
        "link": "https://apix.sn"
    },
    {
        "name": "ADEPME",
        "tags": ["PME", "accompagnement", "diagnostics"],
        "description": "Agence de Développement pour les PME — accompagnement et diagnostics (ne propose plus de formation).",
        "link": "https://adepme.sn"
    },
    {
        "name": "ANPEJ",
        "tags": ["emploi", "jeunes", "formation", "stages"],
        "description": "Agence Nationale pour l'Emploi des Jeunes — formations, stages, dispositifs d’insertion.",
        "link": "https://anpej.sn"
    },
    {
        "name": "ONFP — Office National de Formation Professionnelle",
        "tags": ["formation", "certification", "apprentissage", "professionnelle"],
        "description": "Programmes de formation professionnelle, certifications, apprentissage technique et reconversion.",
        "link": "https://onfp.sn"
    },
    {
        "name": "CBAO / Attijariwafa — Daaray Jàmbaar Yi",
        "tags": ["mentorat", "formation", "financement", "réseau"],
        "description": "Centre d’accompagnement avec mentorat pro, formations et facilitation d’accès au financement.",
        "link": "https://cbao.sn"
    },
    {
        "name": "Bourse Nationale de l’Emploi",
        "tags": ["emploi", "plateforme", "jeunes"],
        "description": "Plateforme d’offres d’emploi et d’opportunités pour les jeunes.",
        "link": "https://bne.sn"
    },
]

# Configuration de la page
st.set_page_config(
    page_title=tr('app_title'),
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)
# Ancre en haut de page pour permettre le lien de remontée
st.markdown("<a id='haut-de-page'></a>", unsafe_allow_html=True)

# CSS personnalisé pour meilleur design
st.markdown("""
<style>
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 20px;
        border-radius: 10px;
        text-align: center;
    }
    .success-card {
        background: #d4edda;
        border-left: 5px solid #28a745;
        padding: 15px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .warning-card {
        background: #fff3cd;
        border-left: 5px solid #ffc107;
        padding: 15px;
        border-radius: 5px;
        margin: 10px 0;
    }
    /* Apparence compacte des boutons numériques */
    .stButton > button {
        min-height: 40px;
        border-radius: 8px;
        font-weight: 600;
    }
    /* Couleur verte pour les boutons primaires (rubriques complétées) */
    [data-testid="baseButton-primary"] {
        background-color: #2ecc71 !important; /* vert */
        color: #ffffff !important;
        border: 1px solid #27ae60 !important;
    }
    [data-testid="baseButton-primary"]:hover {
        background-color: #27ae60 !important;
        border-color: #1e8449 !important;
    }
    /* Style clair pour les boutons secondaires */
    [data-testid="baseButton-secondary"] {
        background-color: #ffffff !important;
        color: #333333 !important;
        border: 1px solid #dddddd !important;
    }
</style>
""", unsafe_allow_html=True)

# Initialisation du client d'analyse
@st.cache_resource
def init_analysis_client(api_key: str | None):
    if not api_key:
        return None
    return OpenAI(
        api_key=api_key,
        base_url="https://api.deepseek.com"
    )

client = None

# Fonctions utilitaires pour la gestion des compétences
def is_competence_completed(competence):
    """Vérifie si une rubrique est complétée"""
    for i in range(len(COMPETENCES[competence]["questions"])):
        selected_key = f"{competence}_{i}"
        if st.session_state.get(selected_key) is None:
            return False
    return True

def all_competences_completed():
    """Vérifie si toutes les rubriques sont complétées"""
    return all(is_competence_completed(comp) for comp in COMPETENCES.keys())

def next_uncompleted_competence(current_comp):
    """Obtient la prochaine rubrique non complétée"""
    names = list(COMPETENCES.keys())
    start = names.index(current_comp) + 1 if current_comp in names else 0
    for offset in range(len(names)):
        comp = names[(start + offset) % len(names)]
        if not is_competence_completed(comp):
            return comp
    return None

def tous_formulaires_remplis(nom, secteur, experience, scores):
    """Vérifie si tous les formulaires sont remplis (nom optionnel)"""
    # Vérifier les informations personnelles (nom optionnel)
    info_complete = all([
        secteur != "Sélectionnez votre secteur",  # Secteur sélectionné
        experience != "Sélectionnez votre niveau"  # Expérience sélectionnée
    ])
    
    # Vérifier que TOUTES les questions de TOUTES les compétences sont répondues
    competences_complete = all_competences_completed()
    
    return info_complete and competences_complete

# Fonction pour générer des recommandations avec streaming
def generate_recommendations_stream(prompt, temperature=0.7):
    # Lecture de la clé API depuis secrets.toml ou variable d'environnement
    api_key = st.secrets.get("deepseek_api_key") or os.environ.get("DEEPSEEK_API_KEY")
    local_client = init_analysis_client(api_key)
    if local_client is None:
        st.warning("Clé API non configurée correctement.")
        return ""
    try:
        stream = local_client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "Tu es un expert en entrepreneuriat et en développement des compétences entrepreneuriales au Sénégal. Tu fournis des analyses précises et des recommandations personnalisées."},
                {"role": "system", "content": get_lang_directive()},
                {"role": "user", "content": prompt}
            ],
            temperature=temperature,
            stream=True
        )
        response_text = ""
        placeholder = st.empty()
        for chunk in stream:
            if chunk.choices[0].delta.content:
                response_text += chunk.choices[0].delta.content
                placeholder.markdown(response_text)
        return response_text
    except Exception as e:
        st.error(f"Erreur lors de la génération des recommandations: {str(e)}")
        return ""

# Chat Coach Fatouma (restriction au domaine entrepreneuriat)
def Fatouma_chat_stream(chat_history, temperature=0.7):
    api_key = st.secrets.get("deepseek_api_key") or os.environ.get("DEEPSEEK_API_KEY")
    local_client = init_analysis_client(api_key)
    if local_client is None:
        st.warning("Clé API non configurée correctement.")
        return ""
    system_persona = {
        "role": "system",
        "content": (
            "Tu es Fatouma, Coach en entrepreneuriat au Sénégal. "
            "Tu réponds uniquement aux questions liées à l'entrepreneuriat: création, gestion, financement, marketing, stratégie, "
            "opérations, leadership, juridique, fiscalité, et ressources locales. "
            "Si une question est hors de ce domaine, réponds seulement: "
            "\"Je suis Coach en entrepreneuriat. Reformule ta question dans ce domaine.\" "
            "Sois claire, concrète et adaptée au contexte sénégalais. "
            "Si le profil de l'utilisateur est disponible, base tes conseils dessus."
        ),
    }
    # Injecter le contexte du profil si disponible, sinon inviter à compléter l'évaluation
    messages = [system_persona, {"role": "system", "content": get_lang_directive()}]
    if st.session_state.get('profil_calcule', False):
        scores = st.session_state.get('scores', {})
        nom = st.session_state.get('nom', 'Non renseigné')
        age = st.session_state.get('age', 'Non renseigné')
        secteur = st.session_state.get('secteur', 'Non spécifié')
        experience = st.session_state.get('experience', 'Non spécifiée')
        try:
            profil, _, _, moyenne = calculer_profil(scores)
        except Exception:
            profil, moyenne = "Non déterminé", 0.0
        contexte_profil = (
            f"Contexte utilisateur connu:\n"
            f"- Nom: {nom}\n- Âge: {age}\n- Secteur: {secteur}\n- Expérience: {experience}\n"
            f"- Profil: {profil}\n- Score global moyen: {moyenne:.2f}/5\n\n"
            f"Scores par compétence:\n" +
            "\n".join([f"- {comp}: {score:.2f}/5" for comp, score in scores.items()])
        )
        messages.append({"role": "system", "content": contexte_profil})
    else:
        messages.append({
            "role": "system",
            "content": (
                "Le profil n'est pas encore rempli. Réponds à la question, puis invite poliment l'utilisateur à compléter l'onglet \"Évaluation\" afin d'obtenir des conseils plus personnalisés."
            )
        })
    messages += chat_history
    try:
        stream = local_client.chat.completions.create(
            model="deepseek-chat",
            messages=messages,
            temperature=temperature,
            stream=True,
        )
        response_text = ""
        placeholder = st.empty()
        for chunk in stream:
            if chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                response_text += chunk.choices[0].delta.content
                placeholder.markdown(response_text)
        return response_text
    except Exception as e:
        st.error(f"Erreur lors du chat avec Fatouma: {str(e)}")
        return ""

# Définition des compétences
COMPETENCES = {
    "Leadership": {
        "questions": [
            "Je prends facilement l'initiative dans un groupe",
            "Je sais motiver et inspirer les autres",
            "Je communique ma vision de façon claire et convaincante",
            "Je sais prendre des décisions difficiles",
            "Je responsabilise mon équipe et favorise l'autonomie",
            "Je favorise la collaboration et résous les conflits efficacement"
        ]
    },
    "Gestion & Délégation": {
        "questions": [
            "Je délègue facilement les tâches à mon équipe",
            "Je fais confiance aux autres pour accomplir des tâches importantes",
            "Je sais organiser et planifier efficacement",
            "Je suis capable de suivre plusieurs projets simultanément",
            "Je définis clairement les priorités et les échéances",
            "Je mets en place des processus pour suivre l’avancement et la qualité"
        ]
    },
    "Créativité & Innovation": {
        "questions": [
            "Je génère facilement des idées nouvelles",
            "J'aime expérimenter de nouvelles approches",
            "Je remets en question le statu quo",
            "Je suis capable d'identifier des opportunités uniques",
            "Je transforme des idées en solutions concrètes",
            "J’observe le marché et j’adapte rapidement mes idées"
        ]
    },
    "Réseautage & Relations": {
        "questions": [
            "Je construis facilement des relations professionnelles",
            "Je maintiens un réseau actif de contacts",
            "Je sais utiliser mon réseau pour atteindre mes objectifs",
            "Je participe activement dans diverses communautés",
            "Je sais entretenir des relations dans la durée",
            "Je crée des partenariats stratégiques bénéfiques aux deux parties"
        ]
    },
    "Résilience & Persévérance": {
        "questions": [
            "Je persiste face aux difficultés",
            "Je maintiens mon focus sur mes objectifs à long terme",
            "Je me relève rapidement après un échec",
            "Je reste positif dans l'adversité",
            "Je garde mon sang-froid sous pression",
            "J’adapte mon plan d’action face aux imprévus sans perdre de vue mes objectifs"
        ]
    },
    "Gestion Financière": {
        "questions": [
            "Je comprends les états financiers de base",
            "Je sais gérer un budget efficacement",
            "Je suis capable d'identifier des sources de financement",
            "Je prends des décisions financières éclairées",
            "Je planifie les flux de trésorerie à moyen terme",
            "Je suis capable de fixer des prix rentables et compétitifs"
        ]
    }
}

def calculer_profil(scores):
    moyenne = sum(scores.values()) / len(scores)
    
    profils = [
        (4.0, "Profil Excellence", "Entrepreneur avec des compétences très développées", "#2E7D32"),
        (3.5, "Profil Avancé", "Entrepreneur expérimenté avec quelques axes d'amélioration", "#558B2F"),
        (3.0, "Profil Intermédiaire", "Entrepreneur en développement avec un potentiel significatif", "#F9A825"),
        (2.5, "Profil Émergent", "Entrepreneur débutant nécessitant un accompagnement ciblé", "#EF6C00"),
        (0, "Profil Débutant", "Entrepreneur ayant besoin d'un accompagnement complet", "#C62828")
    ]
    
    for seuil, profil, desc, couleur in profils:
        if moyenne >= seuil:
            return profil, desc, couleur, moyenne
    
    return profils[-1][1], profils[-1][2], profils[-1][3], moyenne

def creer_diagramme_radar(scores):
    """Crée un beau diagramme radar avec Plotly"""
    categories = list(scores.keys())
    valeurs = list(scores.values())
    
    fig = go.Figure()
    
    fig.add_trace(go.Scatterpolar(
        r=valeurs,
        theta=categories,
        fill='toself',
        name=tr('radar_trace_name'),
        line=dict(color='rgba(102, 126, 234, 0.8)', width=1.5),
        fillcolor='rgba(102, 126, 234, 0.35)',
        hovertemplate=f'<b>%{{theta}}</b><br>{tr("score_label")}: %{{r:.2f}}/5<extra></extra>'
    ))

    # Supprime la ligne horizontale au milieu pour éviter de cacher des libellés
    # fig.add_hline(y=3.0, line_dash="dash", line_color="gray", opacity=0.5)

    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 5],
                showline=False,
                gridcolor='rgba(0,0,0,0.1)',
                gridwidth=0.6,
                tickfont=dict(size=11)
            ),
            angularaxis=dict(
                rotation=90,
                direction='clockwise',
                tickfont=dict(size=14, color='#2c3e50')
            )
        ),
        showlegend=True,
        height=500,
        font=dict(family="Arial, sans-serif", size=12),
        margin=dict(l=50, r=50, t=50, b=50)
    )
    
    return fig

# Interface principale
st.title("🚀 " + tr('app_title'))
st.markdown("### " + tr('app_tagline'))

# Sidebar pour les informations
with st.sidebar:
    st.header(tr('sidebar_info'))
    nom = st.text_input(tr('sidebar_name'), key="nom_input")
    age = st.number_input(tr('sidebar_age'), min_value=18, max_value=100, value=30)
    secteur_options = [
        "Agriculture", "Commerce", "Services", "Technologie",
        "Artisanat", "Transport", "Éducation", "Santé"
    ]
    # Labels Wolof pour les secteurs (affichage), valeurs internes restent en Français
    SECTOR_LABELS = {
        'Français': {
            "Agriculture": "Agriculture",
            "Commerce": "Commerce",
            "Services": "Services",
            "Technologie": "Technologie",
            "Artisanat": "Artisanat",
            "Transport": "Transport",
            "Éducation": "Éducation",
            "Santé": "Santé",
            "Autre (personnalisé)": "Autre (personnalisé)",
        },
        'Wolof': {
            "Agriculture": "Naat",
            "Commerce": "Njënd",
            "Services": "Sarwiis",
            "Technologie": "Teknoloosi",
            "Artisanat": "Artisanaa",
            "Transport": "Transpór",
            "Éducation": "Jàng",
            "Santé": "Wér‑góor",
            "Autre (personnalisé)": "Beneen (sa bopp)",
        }
    }
    def tr_sector(opt: str) -> str:
        lang = st.session_state.get('app_lang', 'Français')
        return SECTOR_LABELS.get(lang, {}).get(opt, opt)

    secteur_choice = st.selectbox(
        tr('sidebar_sector'),
        secteur_options + ["Autre (personnalisé)"],
        key="secteur_select",
        format_func=tr_sector,
    )
    if secteur_choice == "Autre (personnalisé)":
        secteur = st.text_input(
            tr('sidebar_sector_custom'),
            key="secteur_custom",
            placeholder=tr('sidebar_sector_placeholder'),
        )
    else:
        secteur = secteur_choice

    # Labels Wolof pour l'expérience (affichage), valeurs internes restent en Français
    EXPERIENCE_OPTIONS = [
        "Aucune", "Moins de 1 an", "1-3 ans", "3-5 ans", "Plus de 5 ans"
    ]
    EXPERIENCE_LABELS = {
        'Français': {
            "Aucune": "Aucune",
            "Moins de 1 an": "Moins de 1 an",
            "1-3 ans": "1-3 ans",
            "3-5 ans": "3-5 ans",
            "Plus de 5 ans": "Plus de 5 ans",
        },
        'Wolof': {
            "Aucune": "Dara",
            "Moins de 1 an": "Suul 1 at",
            "1-3 ans": "1–3 at",
            "3-5 ans": "3–5 at",
            "Plus de 5 ans": "Lu ëpp 5 at",
        }
    }
    def tr_experience(opt: str) -> str:
        lang = st.session_state.get('app_lang', 'Français')
        return EXPERIENCE_LABELS.get(lang, {}).get(opt, opt)

    experience = st.selectbox(tr('sidebar_experience'), EXPERIENCE_OPTIONS, format_func=tr_experience)
    st.selectbox(tr('sidebar_language'), ["Français", "Wolof"], index=0, key="app_lang")
    # (Champ clé API supprimé)
    
    # Signature
    st.markdown("---")
    st.markdown(f"""
    <div style='text-align: center; font-size: 0.8em; color: #666;'>
        <p><strong>{tr('footer_credit_by')}</strong><br>
        {tr('footer_rights')}<br>
        {tr('footer_phone_label')} 77 359 15 09</p>
    </div>
    """, unsafe_allow_html=True)


# CSS personnalisé pour améliorer la visibilité des onglets
st.markdown("""
<style>
/* Amélioration des onglets */
.stTabs [data-baseweb="tab-list"] {
    gap: 8px;
    background-color: #f8f9fa;
    padding: 10px;
    border-radius: 15px;
    margin-bottom: 20px;
    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
}

.stTabs [data-baseweb="tab"] {
    height: 60px;
    padding: 15px 25px;
    background-color: white;
    border-radius: 12px;
    border: 2px solid #e9ecef;
    font-weight: 600;
    font-size: 16px;
    transition: all 0.3s ease;
    box-shadow: 0 2px 5px rgba(0,0,0,0.05);
}

.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #007bff, #0056b3) !important;
    color: white !important;
    border-color: #007bff !important;
    transform: translateY(-2px);
    box-shadow: 0 4px 15px rgba(0,123,255,0.3) !important;
}

.stTabs [data-baseweb="tab"]:hover {
    background-color: #f8f9fa;
    border-color: #007bff;
    transform: translateY(-1px);
    box-shadow: 0 3px 10px rgba(0,0,0,0.1);
}

/* Indicateurs de progression sur les onglets */
.tab-indicator {
    display: inline-block;
    width: 12px;
    height: 12px;
    border-radius: 50%;
    margin-left: 8px;
    vertical-align: middle;
}

.tab-completed {
    background-color: #28a745;
    animation: pulse 2s infinite;
}

.tab-available {
    background-color: #ffc107;
}

.tab-locked {
    background-color: #6c757d;
}

@keyframes pulse {
    0% { opacity: 1; }
    50% { opacity: 0.5; }
    100% { opacity: 1; }
}

/* Amélioration des messages de navigation */
.navigation-hint {
    background: linear-gradient(135deg, #e3f2fd, #bbdefb);
    border: 2px solid #2196f3;
    border-radius: 15px;
    padding: 20px;
    margin: 20px 0;
    text-align: center;
    font-weight: 600;
    color: #1565c0;
    box-shadow: 0 4px 10px rgba(33,150,243,0.2);
    animation: glow 3s ease-in-out infinite alternate;
}

@keyframes glow {
    from { box-shadow: 0 4px 10px rgba(33,150,243,0.2); }
    to { box-shadow: 0 6px 20px rgba(33,150,243,0.4); }
}

.next-step-button {
    background: linear-gradient(135deg, #28a745, #20c997) !important;
    color: white !important;
    border: none !important;
    border-radius: 25px !important;
    padding: 15px 30px !important;
    font-weight: 600 !important;
    font-size: 16px !important;
    box-shadow: 0 4px 15px rgba(40,167,69,0.3) !important;
    transition: all 0.3s ease !important;
}

.next-step-button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 20px rgba(40,167,69,0.4) !important;
}
</style>
""", unsafe_allow_html=True)

# Déterminer l'état des onglets
evaluation_complete = st.session_state.get('profil_calcule', False)
results_available = evaluation_complete
recommendations_available = evaluation_complete

# Créer les labels des onglets avec indicateurs (localisés)
tab1_label = "📝 " + tr('tab_eval')
if evaluation_complete:
    tab1_label += " ✅"

tab2_label = "📊 " + tr('tab_results')
if results_available:
    tab2_label += " ✅"
elif evaluation_complete:
    tab2_label += " 🔓"
else:
    tab2_label += " 🔒"

tab3_label = "💡 " + tr('tab_reco')
if recommendations_available:
    tab3_label += " ✅"
elif evaluation_complete:
    tab3_label += " 🔓"
else:
    tab3_label += " 🔒"

# Tabs pour l'interface avec labels améliorés
tab4_label = "👩🏾‍💼 " + tr('tab_adja')
tab1, tab2, tab3, tab4 = st.tabs([tab1_label, tab2_label, tab3_label, tab4_label])

with tab1:
    st.header(tr('tab1_header'))
    st.markdown("*" + tr('tab1_instruction') + "*")
    # Barre de progression globale
    total_questions = sum(len(data["questions"]) for data in COMPETENCES.values())
    answered_questions = sum(1 for comp, data in COMPETENCES.items() for i in range(len(data["questions"])) if st.session_state.get(f"{comp}_{i}") is not None)
    progress_ratio = (answered_questions / total_questions) if total_questions else 0
    st.progress(progress_ratio, text=f"{tr('progress_global')}: {progress_ratio*100:.0f}%")
    
    # Interface avec rubriques cliquables
    if 'selected_competence' not in st.session_state:
        st.session_state.selected_competence = None
    
    # Affichage des rubriques en ligne
    st.subheader("📋 " + tr('tab1_rubriques'))
    

    
    # Créer des colonnes pour les boutons de rubriques
    competence_names = list(COMPETENCES.keys())
    cols = st.columns(3)  # 3 colonnes pour 6 rubriques
    
    for i, competence in enumerate(competence_names):
        with cols[i % 3]:
            # Vérifier si la rubrique est complétée
            is_completed = is_competence_completed(competence)
            
            # Déterminer si cette rubrique est sélectionnée
            is_selected = st.session_state.selected_competence == competence
            
            # Déterminer le style et le texte du bouton
            # Vert (primary) uniquement si complétée, sinon neutre (secondary)
            if is_completed:
                button_text = f"{tr_comp(competence)}"
                button_style = "primary"
            else:
                button_text = f"🎯 {tr_comp(competence)}"
                button_style = "secondary"
            
            if st.button(
                button_text,
                key=f"rubrique_{competence}",
                type=button_style,
                use_container_width=True
            ):
                st.session_state.selected_competence = competence
                try:
                    st.rerun()
                except Exception:
                    st.experimental_rerun()
    
    # Affichage des questions pour la rubrique sélectionnée
    if st.session_state.selected_competence:
        st.markdown("---")
        selected_comp = st.session_state.selected_competence
        
        # Titre de la rubrique sélectionnée (plus compact)
        st.subheader(f"🎯 {tr_comp(selected_comp)}")
        
        # Questions de la rubrique sélectionnée (format compact)
        with st.container():
            # Afficher toutes les questions en format compact
            for i, question in enumerate(COMPETENCES[selected_comp]["questions"]):
                selected_key = f"{selected_comp}_{i}"
                selected = st.session_state.get(selected_key)
                
                # Question et boutons sur la même ligne
                col_question, col_buttons = st.columns([3, 2])
                
                with col_question:
                    st.write(f"**{i+1}.** {tr_question(selected_comp, i, question)}")
                
                with col_buttons:
                    cols_nums = st.columns(5)
                    for val in range(1, 6):
                        with cols_nums[val - 1]:
                            if st.button(
                                str(val),
                                key=f"{selected_key}_btn_{val}",
                                type="primary" if selected == val else "secondary",
                                use_container_width=True,
                            ):
                                st.session_state[selected_key] = val
                                try:
                                    st.rerun()
                                except Exception:
                                    st.experimental_rerun()
                
                # Affichage compact du statut
                if selected is not None:
                    st.caption(f"✅ {selected}/5")
                else:
                    st.caption("⏳ " + tr('to_evaluate'))
    
    else:
        st.info(tr('click_rubrique_hint'))
    
    # Calcul des scores pour toutes les compétences
    scores = {}
    
    for competence, data in COMPETENCES.items():
        questions_scores = []
        for i, question in enumerate(data["questions"]):
            selected_key = f"{competence}_{i}"
            selected = st.session_state.get(selected_key)
            if selected is not None:
                questions_scores.append(selected)
        
        # Moyenne par compétence (0.0 si aucune réponse)
        scores[competence] = (sum(questions_scores) / len(questions_scores)) if questions_scores else 0.0
    # Vérification automatique et calcul du profil
    formulaires_remplis = tous_formulaires_remplis(nom, secteur, experience, scores)
    
    if formulaires_remplis:
        # Mise à jour automatique du profil à chaque modification
        scores_changed = st.session_state.get('scores') != scores
        info_changed = (st.session_state.get('nom') != nom or 
                       st.session_state.get('age') != age or 
                       st.session_state.get('secteur') != secteur or 
                       st.session_state.get('experience') != experience)
        
        # Recalculer si c'est la première fois ou si quelque chose a changé
        if not st.session_state.get('profil_calcule', False) or scores_changed or info_changed:
            st.session_state.scores = scores
            st.session_state.profil_calcule = True
            st.session_state.nom = nom
            st.session_state.age = age
            st.session_state.secteur = secteur
            st.session_state.experience = experience
            
            if not st.session_state.get('profil_calcule', False):
                st.success("🎉 Profil calculé automatiquement ! Consultez l'onglet 'Résultats' pour voir vos graphiques et recommandations.")
            else:
                st.info("🔄 Profil mis à jour automatiquement suite à vos modifications.")
            
            try:
                st.rerun()
            except Exception:
                st.experimental_rerun()
        
        # Affichage du statut de completion
        col1, col2, col3 = st.columns(3)
        with col2:
            st.markdown(f"""
            <div style="
                background: linear-gradient(135deg, #28a745, #20c997);
                color: white;
                padding: 15px;
                border-radius: 10px;
                text-align: center;
                font-weight: bold;
                box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
            ">
                {tr('profile_calculated')}
            </div>
            """, unsafe_allow_html=True)
    else:
        # Affichage du bouton désactivé avec indication du progrès
        col1, col2, col3 = st.columns(3)
        with col2:
            # Calcul du pourcentage de completion
            total_questions = sum(len(data["questions"]) for data in COMPETENCES.values())
            answered_questions = sum(1 for competence, data in COMPETENCES.items() 
                                   for i in range(len(data["questions"])) 
                                   if st.session_state.get(f"{competence}_{i}") is not None)
            
            # Calcul du pourcentage de progression basé sur les questions répondues
            progress_text = tr('questions_answered').format(answered=answered_questions, total=total_questions)
            progress_percent = (answered_questions / total_questions) * 100 if total_questions > 0 else 0
            
            st.markdown(f"""
            <div style="
                background: #f8f9fa;
                border: 2px solid #dee2e6;
                color: #6c757d;
                padding: 15px;
                border-radius: 10px;
                text-align: center;
                font-weight: bold;
            ">
                🔄 {tr('complete_info')}<br>
                <small>{progress_text}</small><br>
                <div style="background: #e9ecef; border-radius: 10px; height: 8px; margin: 10px 0;">
                    <div style="background: #007bff; height: 100%; width: {progress_percent}%; border-radius: 10px;"></div>
                </div>
            </div>
            """, unsafe_allow_html=True)
    with col3:
        # Afficher le bouton "Rubrique suivante" seulement si une rubrique est sélectionnée ET qu'il reste des rubriques incomplètes
        if st.session_state.get('selected_competence') and not all_competences_completed():
            if st.button(tr('next_rubrique_button'), key="btn_next_rubrique", type="secondary", use_container_width=True):
                target = next_uncompleted_competence(st.session_state.get('selected_competence'))
                if target:
                    st.session_state.selected_competence = target
                    try:
                        st.rerun()
                    except Exception:
                        st.experimental_rerun()
    # Résumé rapide directement sous le bouton pour éviter de remonter
    if st.session_state.get('profil_calcule'):
        st.markdown("### " + tr('resume_rapide'))
        profil, description, couleur, moyenne = calculer_profil(st.session_state.scores)
        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric(tr('score_global'), f"{moyenne:.2f}/5")
        with c2:
            st.markdown(f"<div style='text-align: center; padding: 12px; background: {couleur}22; border-radius: 8px; border-left: 4px solid {couleur}'><b style='color: {couleur}'>{profil}</b></div>", unsafe_allow_html=True)
        with c3:
            pf = sum(1 for s in st.session_state.scores.values() if s >= 4.0)
            st.metric(tr('points_forts'), f"{pf}/{len(st.session_state.scores)}")
        # Mini-diagramme
        st.plotly_chart(creer_diagramme_radar(st.session_state.scores), use_container_width=True, key="radar_summary_tab1")
        # Info + lien de remontée
        cInfo, cBtn = st.columns([3, 1])
        with cInfo:
            st.info(tr('open_results_hint'))
        with cBtn:
            st.markdown(f"""
            <div style="
                background: linear-gradient(135deg, #007bff, #0056b3);
                color: white;
                padding: 8px 12px;
                border-radius: 20px;
                text-align: center;
                box-shadow: 0 2px 8px rgba(0,123,255,0.3);
                transition: all 0.3s ease;
            ">
                <a href='#haut-de-page' style='color: white; text-decoration: none; font-weight: 600;'>
                    {tr('monter')}
                </a>
            </div>
            """, unsafe_allow_html=True)
    
    # Message de navigation pour guider l'utilisateur
    if st.session_state.get('profil_calcule'):
        st.markdown(f"""
        <div class="navigation-hint">
            {tr('nav_success_heading')}<br>
            {tr('nav_results_hint')}<br>
            {tr('nav_reco_hint')}
        </div>
        """, unsafe_allow_html=True)
    elif all_competences_completed():
        st.markdown(f"""
        <div class="navigation-hint">
            {tr('nav_all_completed')}<br>
            {tr('nav_complete_personal_info')}
        </div>
        """, unsafe_allow_html=True)
    else:
        progress_percent = (sum(1 for comp in COMPETENCES.keys() if is_competence_completed(comp)) / len(COMPETENCES)) * 100
        st.markdown(f"""
        <div class="navigation-hint">
            {tr('nav_progress_label').format(percent=f"{progress_percent:.0f}")}<br>
            {tr('nav_continue_eval')}
        </div>
        """, unsafe_allow_html=True)

with tab2:
    if 'profil_calcule' in st.session_state and st.session_state.profil_calcule:
        scores = st.session_state.scores
        nom = st.session_state.get('nom', tr('non_renseigne'))
        
        # Saisie optionnelle du nom de l'entreprise
        st.text_input(tr('company_name_label'), key="entreprise_tab3")
        
        st.header(tr('results_header'))
        
        profil, description, couleur, moyenne = calculer_profil(scores)
        
        # Cartes de métriques
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric(tr('score_global'), f"{moyenne:.2f}/5", delta=tr('delta_excellent') if moyenne >= 3.5 else tr('delta_to_develop'))
        
        with col2:
            st.markdown(f"<div style='text-align: center; padding: 20px; background: {couleur}22; border-radius: 10px; border-left: 5px solid {couleur}'><h3 style='color: {couleur}; margin: 0'>{profil}</h3></div>", unsafe_allow_html=True)
        
        with col3:
            points_forts = sum(1 for s in scores.values() if s >= 4.0)
            st.metric(tr('points_forts'), f"{points_forts}/{len(scores)}")
        
        st.info(f"📌 {description}")
        
        # Diagramme radar amélioré
        st.subheader(tr('radar_map_title'))
        fig_radar = creer_diagramme_radar(scores)
        st.plotly_chart(fig_radar, use_container_width=True, key="radar_full_tab3")
        
        # 🔥 Heatmap des compétences
        st.markdown("### " + tr('heatmap_comp_title'))
        heatmap_fig = go.Figure(data=go.Heatmap(
            z=[list(scores.values())],
            x=list(scores.keys()),
            y=[tr('score_label')],
            colorscale='YlOrRd', zmin=0, zmax=5, showscale=True
        ))
        heatmap_fig.update_layout(height=180, margin=dict(l=10, r=10, t=10, b=10))
        st.plotly_chart(heatmap_fig, use_container_width=True)

        # 🏅 Badges
        st.markdown("### 🏅 Badges")
        badges = []
        if moyenne >= 3.5:
            badges.append("Excellence (Score global ≥ 3.5)")
        if all_competences_completed():
            badges.append("Évaluation complète")
        if badges:
            st.success(" | ".join([f"🏅 {b}" for b in badges]))
        else:
            st.info("Complétez l'évaluation pour débloquer des badges.")

        # 💾 Export CSV des scores
        st.download_button(
            label="💾 Télécharger Scores (CSV)",
            data=make_scores_csv(scores),
            file_name=f"scores_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
            key="dl_scores_csv"
        )
        
        # Bouton unique pleine largeur pour déclencher les recommandations sommaires
        if st.button("💡 Recommandations Sommaires - Cliquez ici !", type="primary", use_container_width=True, key="reco_sommaire_duplicate", help="Obtenez des recommandations personnalisées basées sur votre profil"):
                st.subheader("💡 Recommandations Personnalisées")
                with st.spinner("Génération des recommandations en cours..."):
                    # Préparer le contexte pour l'analyse
                    contexte_sommaire = f"""
 Profil entrepreneur: {profil}
 Score global: {sum(scores.values()) / len(scores):.2f}/5
 Secteur: {st.session_state.get('secteur', 'Non spécifié')}
 Expérience: {st.session_state.get('experience', 'Non spécifiée')}
 
 Scores détaillés:
 """
                    for comp, score in scores.items():
                        contexte_sommaire += f"- {comp}: {score:.2f}/5\n"
                    prompt_sommaire = f"""{contexte_sommaire}

En tant qu'expert en entrepreneuriat au Sénégal, fournis 3-4 recommandations courtes et concrètes (maximum 150 mots) pour cet entrepreneur basées sur son profil.

Focus sur:
1. Les 2 compétences les plus faibles à améliorer en priorité
2. Une action concrète à mettre en place dans les 30 prochains jours
3. Une ressource ou contact utile au Sénégal

Sois direct, actionnable et adapté au contexte sénégalais."""
                    reponse_sommaire = generate_recommendations_stream(prompt_sommaire)
                    st.session_state['reco_sommaire_text'] = reponse_sommaire
                    st.success("✅ Recommandations sommaires enregistrées pour le rapport.")
        
        # Analyse détaillée
        st.subheader("📈 Analyse Détaillée")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### ✅ " + tr('points_forts'))
            points_forts_list = sorted(scores.items(), key=lambda x: x[1], reverse=True)[:3]
            for i, (comp, score) in enumerate(points_forts_list, 1):
                st.markdown(f"""
                <div class='success-card'>
                    <b style='font-size: 1.1em;'>#{i} {comp}</b><br>
                    Score: {score:.2f}/5 ⭐
                </div>
                """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("### ⚠️ " + tr('axes_amelioration'))
            axes_amelioration = sorted(scores.items(), key=lambda x: x[1])[:3]
            for i, (comp, score) in enumerate(axes_amelioration, 1):
                st.markdown(f"""
                <div class='warning-card'>
                    <b style='font-size: 1.1em;'>#{i} {comp}</b><br>
                    Score: {score:.2f}/5 📈
                </div>
                """, unsafe_allow_html=True)
        
        # Grille de développement
        st.subheader("🎯 Positionnement dans la Grille de Développement")
        
        # Déterminer le niveau actuel basé sur le score moyen
        score_moyen = sum(scores.values()) / len(scores)
        
        if score_moyen < 2.5:
            niveau_actuel = "Débutant"
            couleur_niveau = "#ff4b4b"
            emoji_niveau = "🔴"
        elif score_moyen < 3.0:
            niveau_actuel = "Émergent"
            couleur_niveau = "#ff8c00"
            emoji_niveau = "🟠"
        elif score_moyen < 3.5:
            niveau_actuel = "Intermédiaire"
            couleur_niveau = "#ffd700"
            emoji_niveau = "🟡"
        elif score_moyen < 4.0:
            niveau_actuel = "Avancé"
            couleur_niveau = "#32cd32"
            emoji_niveau = "🟢"
        else:
            niveau_actuel = "Excellence"
            couleur_niveau = "#00ff00"
            emoji_niveau = "🌟"
        
        # Affichage visuel du niveau actuel
        st.markdown(f"""
        <div style="
            background: linear-gradient(135deg, {couleur_niveau}20, {couleur_niveau}10);
            border-left: 5px solid {couleur_niveau};
            padding: 20px;
            border-radius: 10px;
            margin: 20px 0;
            text-align: center;
        ">
            <h3 style="color: {couleur_niveau}; margin: 0;">
                {emoji_niveau} Votre Niveau Actuel : {niveau_actuel}
            </h3>
            <p style="font-size: 18px; margin: 10px 0; color: #333;">
                Score Global : <strong>{score_moyen:.2f}/5.0</strong>
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # Grille de progression visuelle
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("### 📊 Grille de Progression")
            
            # Créer une visualisation de la grille avec barres de progression
            niveaux = ["Débutant", "Émergent", "Intermédiaire", "Avancé", "Excellence"]
            scores_min = [0, 2.5, 3.0, 3.5, 4.0]
            scores_max = [2.5, 3.0, 3.5, 4.0, 5.0]
            couleurs = ["#ff4b4b", "#ff8c00", "#ffd700", "#32cd32", "#00ff00"]
            emojis = ["🔴", "🟠", "🟡", "🟢", "🌟"]
            
            for i, (niveau, score_min, score_max, couleur, emoji) in enumerate(zip(niveaux, scores_min, scores_max, couleurs, emojis)):
                # Calculer le pourcentage de progression dans ce niveau
                if score_moyen >= score_max:
                    progression = 100
                elif score_moyen >= score_min:
                    progression = ((score_moyen - score_min) / (score_max - score_min)) * 100
                else:
                    progression = 0
                
                # Indicateur si c'est le niveau actuel
                indicateur = (" ← " + tr('vous_etes_ici')) if niveau == niveau_actuel else ""
                
                st.markdown(f"""
                <div style="margin: 10px 0;">
                    <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 5px;">
                        <span style="font-weight: bold;">{emoji} {niveau}{indicateur}</span>
                        <span style="font-size: 12px; color: #666;">{score_min:.1f} - {score_max:.1f}</span>
                    </div>
                    <div style="background-color: #f0f0f0; border-radius: 10px; height: 20px; overflow: hidden;">
                        <div style="
                            background-color: {couleur};
                            height: 100%;
                            width: {progression}%;
                            border-radius: 10px;
                            transition: width 0.3s ease;
                        "></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("### " + tr('actions_recommandees'))
            
            actions = {
                "Débutant": "📚 Formation complète",
                "Émergent": "🎯 Accompagnement ciblé", 
                "Intermédiaire": "👥 Mentorat",
                "Avancé": "⚡ Perfectionnement",
                "Excellence": "🌟 Partage d'expertise"
            }
            
            st.markdown(f"""
            <div style="
                background: {couleur_niveau}20;
                border: 2px solid {couleur_niveau};
                border-radius: 10px;
                padding: 15px;
                text-align: center;
                margin-top: 20px;
            ">
                <h4 style="color: {couleur_niveau}; margin: 0;">
                    {actions[niveau_actuel]}
                </h4>
            </div>
            """, unsafe_allow_html=True)
        
        # Téléchargement du rapport
        rapport = {
            "nom": st.session_state.get('nom', 'N/A'),
            "entreprise": st.session_state.get('entreprise', 'N/A'),
            "age": st.session_state.get('age', 'N/A'),
            "secteur": st.session_state.get('secteur', 'N/A'),
            "experience": st.session_state.get('experience', 'N/A'),
            "profil": profil,
            "description": description,
            "scores": scores,
        }
        
        # Générer un rapport Word avec image du radar

        if st.button("📝 Générer le rapport Word", type="primary", use_container_width=True, key="btn_gen_word_duplicate"):
            buf = io.BytesIO()
            
            doc = Document()
            doc.add_heading("Rapport de Profilage Entrepreneurial", level=1)
            doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M"))
            
            doc.add_heading("Informations", level=2)
            doc.add_paragraph(f"Nom: {rapport['nom']}")
            doc.add_paragraph(f"Entreprise: {rapport['entreprise']}")
            doc.add_paragraph(f"Âge: {rapport['age']}")
            doc.add_paragraph(f"Secteur: {rapport['secteur']}")
            doc.add_paragraph(f"Expérience: {rapport['experience']}")
            
            doc.add_heading("Synthèse du Profil", level=2)
            doc.add_paragraph(f"Profil: {rapport['profil']}")
            doc.add_paragraph(rapport['description'])
            
            doc.add_heading("Scores par Compétence", level=2)
            for comp, sc in rapport['scores'].items():
                doc.add_paragraph(f"- {comp}: {sc:.2f}/5")
            
            # Note: Image radar supprimée pour éviter les lenteurs de calcul
            doc.add_heading("Cartographie des Compétences", level=2)
            doc.add_paragraph("Consultez l'application pour visualiser le diagramme radar interactif.")
            
            # Inclure les recommandations sommaires seulement si générées
            reco_text = st.session_state.get('reco_sommaire_text')
            if reco_text and reco_text.strip():
                doc.add_heading("Recommandations Sommaires", level=2)
                for line in reco_text.splitlines():
                    doc.add_paragraph(line)
            
            doc.save(buf)
            buf.seek(0)
            
            st.download_button(
                label="💾 Télécharger mon rapport (Word)",
                data=buf.getvalue(),
                file_name=f"rapport_profil_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        
        # Message de navigation vers les recommandations
        st.markdown(f"""
        <div class="navigation-hint">
            {tr('nav_reco_hint')}
        </div>
        """, unsafe_allow_html=True)
        
    else:
        st.info(tr('goto_eval_warning'))

with tab3:
    if 'profil_calcule' in st.session_state and st.session_state.profil_calcule:
        st.header("💡 " + tr('tab_reco'))
        
        scores = st.session_state.scores
        nom = st.session_state.get('nom', 'Non renseigné')
        age = st.session_state.get('age', 30)
        secteur = st.session_state.get('secteur', 'Non spécifié')
        experience = st.session_state.get('experience', 'Non spécifiée')
        profil, description, _, _ = calculer_profil(scores)
        
        st.markdown("### Analyse approfondie et recommandations personnalisées")
        
        # Préparer le contexte
        contexte = f"""
Contexte de l'entrepreneur:
- Nom: {nom}
- Âge: {age}
- Secteur: {secteur}
- Expérience: {experience}
- Profil identifié: {profil}

Scores par compétence:
"""
        for comp, score in scores.items():
            contexte += f"- {comp}: {score:.2f}/5\n"
        
        # Boutons pour recommandations avec colonnes
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📚 Plan de Formation Personnalisé", use_container_width=True, key="formation"):
                st.subheader("📚 Plan de Formation Personnalisé")
                with st.spinner("Génération en cours..."):
                    prompt = f"""{contexte}

En tant qu'expert en formation entrepreneuriale au Sénégal, propose un plan de formation détaillé et personnalisé pour cet entrepreneur. 
Inclus:
1. Les domaines prioritaires à développer
2. Des formations spécifiques recommandées (disponibles au Sénégal)
3. Un calendrier suggéré sur 6-12 mois
4. Des ressources locales (organisations, programmes, institutions sénégalaises)

RESSOURCES SPÉCIFIQUES À MENTIONNER SI PERTINENTES :
        - Daaray Jàmbaar Yi (CBAO Groupe Attijariwafa bank) : centre d'accompagnement offrant formations entrepreneuriales, coaching personnalisé, aide au montage de projets, business plans, et facilitation d'accès au financement. Idéal pour initiation à l'entrepreneuriat, modules spécialisés (business model, gestion d'entreprise, éducation financière) et accompagnement des TPME/PME.
        - ONFP — Office National de Formation Professionnelle : programmes de formation professionnelle, certifications, apprentissage technique et reconversion, adaptés au développement des compétences métiers.
        NOTE FORMATION : ADEPME n’offre plus de formation ; ne pas la recommander pour ce volet."""
                    
                    reponse_formation = generate_recommendations_stream(prompt)
                    
                    # Boutons de téléchargement
                    col_txt, col_word = st.columns(2)
                    with col_txt:
                        st.download_button(
                            label="💾 Télécharger en TXT",
                            data=reponse_formation,
                            file_name=f"plan_formation_{datetime.now().strftime('%Y%m%d')}.txt",
                            mime="text/plain",
                            key="dl_formation_txt"
                        )
                    with col_word:
                        st.download_button(
                            label="📄 Télécharger en Word",
                            data=make_docx("Plan de Formation Personnalisé", reponse_formation),
                            file_name=f"plan_formation_{datetime.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_formation_word"
                        )
        
        with col2:
            if st.button("🎯 Stratégie de Développement", use_container_width=True, key="strategie"):
                st.subheader("🎯 Stratégie de Développement")
                with st.spinner("Génération en cours..."):
                    prompt = f"""{contexte}

En tant qu'expert en développement entrepreneurial, propose une stratégie de développement sur mesure pour cet entrepreneur sénégalais.
Inclus:
1. Des objectifs SMART à court terme (3 mois)
2. Des objectifs à moyen terme (6-12 mois)
3. Des actions concrètes et mesurables
4. Des indicateurs de succès
5. Des opportunités spécifiques au contexte sénégalais

RESSOURCES SPÉCIFIQUES À MENTIONNER SI PERTINENTES :
- Daaray Jàmbaar Yi (CBAO Groupe Attijariwafa bank) : pour coaching personnalisé, mentorat par des professionnels bancaires, conseils pour optimiser l'accès au financement, et networking avec chefs d'entreprise et investisseurs."""
                    
                    reponse_strategie = generate_recommendations_stream(prompt)
                    
                    # Boutons de téléchargement
                    col_txt, col_word = st.columns(2)
                    with col_txt:
                        st.download_button(
                            label="💾 Télécharger en TXT",
                            data=reponse_strategie,
                            file_name=f"strategie_developpement_{datetime.now().strftime('%Y%m%d')}.txt",
                            mime="text/plain",
                            key="dl_strategie_txt"
                        )
                    with col_word:
                        st.download_button(
                            label="📄 Télécharger en Word",
                            data=make_docx("Stratégie de Développement", reponse_strategie),
                            file_name=f"strategie_developpement_{datetime.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_strategie_word"
                        )
        
        col3, col4 = st.columns(2)
        
        with col3:
            if st.button(tr('mentorat_button'), use_container_width=True, key="mentorat"):
                st.subheader(tr('mentorat_button'))
                with st.spinner(tr('generating')):
                    prompt = f"""{contexte}

Recommande un programme de mentorat adapté à cet entrepreneur sénégalais.
Inclus:
1. Le type de mentor idéal (profil, expérience)
2. Les domaines où le mentorat est le plus nécessaire
3. Des programmes de mentorat disponibles au Sénégal
4. Comment tirer le meilleur parti du mentorat
5. Des structures d'accompagnement locales (incubateurs, accélérateurs)

RESSOURCES SPÉCIFIQUES À MENTIONNER SI PERTINENTES :
- Daaray Jàmbaar Yi (CBAO Groupe Attijariwafa bank) : offre mentorat par des professionnels du secteur bancaire et de l'entreprise, suivi individuel des porteurs de projet, sessions de rencontres avec chefs d'entreprise et investisseurs, et plateforme d'échanges entre entrepreneurs."""
                    
                    reponse_mentorat = generate_recommendations_stream(prompt)
                    
                    # Boutons de téléchargement
                    col_txt, col_word = st.columns(2)
                    with col_txt:
                        st.download_button(
                            label=tr('download_txt'),
                            data=reponse_mentorat,
                            file_name=f"recommandations_mentorat_{datetime.now().strftime('%Y%m%d')}.txt",
                            mime="text/plain",
                            key="dl_mentorat_txt"
                        )
                    with col_word:
                        st.download_button(
                            label=tr('download_word'),
                            data=make_docx(tr('doc_title_mentorat'), reponse_mentorat),
                            file_name=f"recommandations_mentorat_{datetime.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_mentorat_word"
                        )
        
        with col4:
            if st.button(tr('financement_button'), use_container_width=True, key="financement"):
                st.subheader(tr('financement_button'))
                with st.spinner("Génération en cours..."):
                    prompt = f"""{contexte}

Identifie les opportunités de financement adaptées à cet entrepreneur sénégalais.
Inclus:
1. Les types de financement recommandés selon son profil
2. Des programmes de financement disponibles au Sénégal
3. Les critères d'éligibilité typiques
4. Comment renforcer sa candidature
5. Des alternatives au financement traditionnel

RESSOURCES SPÉCIFIQUES À MENTIONNER SI PERTINENTES :
- Daaray Jàmbaar Yi (CBAO Groupe Attijariwafa bank) : facilite l'accès au crédit et aux services bancaires, partenariats privilégiés avec la CBAO pour TPME/PME, information sur produits bancaires adaptés aux petites structures, et appui pour monter un dossier de crédit ou de financement adapté."""
                    
                    reponse_financement = generate_recommendations_stream(prompt)
                    
                    # Boutons de téléchargement
                    col_txt, col_word = st.columns(2)
                    with col_txt:
                        st.download_button(
                            label=tr('download_txt'),
                            data=reponse_financement,
                            file_name=f"opportunites_financement_{datetime.now().strftime('%Y%m%d')}.txt",
                            mime="text/plain",
                            key="dl_financement_txt"
                        )
                    with col_word:
                        st.download_button(
                            label=tr('download_word'),
                            data=make_docx(tr('doc_title_financement'), reponse_financement),
                            file_name=f"opportunites_financement_{datetime.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_financement_word"
                        )

        # 🗓️ Plan d'action 90 jours
        st.markdown("### " + tr('plan_action_90_title'))
        col_plan1, col_plan2 = st.columns([2, 1])
        with col_plan1:
            if st.button(tr('plan_action_90_generate'), use_container_width=True, key="plan_90"):
                st.subheader(tr('plan_action_90_title'))
                with st.spinner(tr('generating')):
                    prompt = f"""{contexte}

En tant que conseiller en entrepreneuriat au Sénégal, crée un plan d'action structuré sur 90 jours:
- Semaines 1-4: Actions immédiates (marketing, opérations, finances)
- Semaines 5-8: Consolidation (processus, équipe, partenariats)
- Semaines 9-12: Évaluation et ajustement

Inclure: objectifs mesurables, tâches concrètes, indicateurs de succès, et ressources locales pertinentes.
"""
                    reponse_plan = generate_recommendations_stream(prompt)
                    st.session_state['plan_90_text'] = reponse_plan
        with col_plan2:
            if st.session_state.get('plan_90_text'):
                encoded = urllib.parse.quote(st.session_state['plan_90_text'])
                st.markdown(f"[{tr('share_whatsapp')}](https://wa.me/?text={encoded})")
                st.download_button(
                    label=tr('download_txt'),
                    data=st.session_state['plan_90_text'],
                    file_name=f"plan_90_jours_{datetime.now().strftime('%Y%m%d')}.txt",
                    mime="text/plain",
                    key="dl_plan90_txt"
                )
                st.download_button(
                    label=tr('download_word'),
                    data=make_docx(tr('plan_action_90_title'), st.session_state['plan_90_text']),
                    file_name=f"plan_90_jours_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_plan90_docx"
                )

        # 📚 Ressources Locales (Recherche)
        st.markdown("### " + tr('local_resources_title'))
        query = st.text_input(tr('search_resources_placeholder'), key="search_resources")
        def match_res(r, q):
            q = q.lower()
            return (
                q in r["name"].lower() or
                any(q in t.lower() for t in r["tags"]) or
                q in r["description"].lower()
            )
        filtered = LOCAL_RESOURCES if not query or not query.strip() else [r for r in LOCAL_RESOURCES if match_res(r, query.strip())]
        for r in filtered:
            link = f" [Lien]({r['link']})" if r.get('link') else ""
            tags = ", ".join(r["tags"]) if r.get("tags") else ""
            st.markdown(f"- **{r['name']}** — {r['description']} ({tags}){link}")
        if not filtered:
            st.info(tr('no_resource_match'))
        
        st.markdown("---")
        
        # Analyse complète
        if st.button(tr('analyse_complete_button'), type="primary", use_container_width=True):
            st.subheader(tr('analyse_complete_button'))
            with st.spinner(tr('generating')):
                prompt = f"""{contexte}

En tant qu'expert en entrepreneuriat au Sénégal, fournis une analyse complète et des recommandations globales pour cet entrepreneur.

Structure ton analyse ainsi:

1. **ANALYSE DU PROFIL**
   - Forces principales
   - Faiblesses critiques
   - Opportunités de développement

2. **RECOMMANDATIONS PRIORITAIRES**
   - Top 3 des compétences à développer en urgence
   - Actions concrètes pour chaque compétence
   - Délais recommandés

3. **PLAN D'ACTION 90 JOURS**
   - Semaines 1-4: Actions immédiates
   - Semaines 5-8: Consolidation
   - Semaines 9-12: Évaluation et ajustement

4. **RESSOURCES SPÉCIFIQUES AU SÉNÉGAL**
   - Organisations d'accompagnement
   - Programmes de formation
   - Réseaux d'entrepreneurs
   - Opportunités de financement

RESSOURCE PRIORITAIRE À MENTIONNER :
- Daaray Jàmbaar Yi (CBAO Groupe Attijariwafa bank) : centre d'accompagnement complet offrant formations entrepreneuriales, coaching personnalisé, mentorat par professionnels bancaires, aide au montage de projets et business plans, facilitation d'accès au financement, et networking avec entrepreneurs et investisseurs. Idéal pour tous profils d'entrepreneurs (jeunes porteurs de projet, TPME, PME, femmes entrepreneures).

RESSOURCES FORMATION À PRIVILÉGIER :
- ONFP — Office National de Formation Professionnelle : programmes de formation professionnelle, certifications, apprentissage technique et reconversion, adaptés au développement des compétences métiers.
NOTE FORMATION : ADEPME n’offre plus de formation ; ne pas la recommander pour ce volet.

5. **CONSEILS ADAPTÉS AU SECTEUR** ({secteur})
   - Spécificités du secteur au Sénégal
   - Meilleures pratiques
   - Pièges à éviter

Sois concret, actionnable et adapté au contexte sénégalais."""
                
                reponse = generate_recommendations_stream(prompt)
                
                # Option de téléchargement
                st.download_button(
                    label=tr('download_analysis_complete'),
                    data=reponse,
                    file_name=f"analyse_complete_{datetime.now().strftime('%Y%m%d')}.txt",
                    mime="text/plain"
                )
                st.download_button(
                    label=tr('download_analysis_word'),
                    data=make_docx(tr('doc_title_analyse_complete'), reponse),
                    file_name=f"analyse_complete_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# Footer
with tab4:
    st.header("👩🏾‍💼 " + tr('tab_adja'))
    st.caption(tr('adja_caption'))
    if 'coaching_journal' not in st.session_state:
        st.session_state['coaching_journal'] = []
    if 'Fatouma_chat' not in st.session_state:
        # Message d’accueil selon la langue
        lang = st.session_state.get('app_lang', 'Français')
        if lang == 'Wolof':
            welcome = "Salaamaleekum, Cooc Fatouma laa. Ci entrepreneuriat ci Senegaal laa. Laaj sa laaj bu jëm ci entrepreneuriat."
        else:
            welcome = "Bonjour, je suis Coach Fatouma, spécialisée en entrepreneuriat au Sénégal. Pose ta question liée à l’entrepreneuriat."
        st.session_state['Fatouma_chat'] = [{"role": "assistant", "content": welcome}]
    for msg in st.session_state['Fatouma_chat']:
        st.chat_message(msg["role"]).markdown(msg["content"])
    # Placeholder de saisi selon la langue
    lang = st.session_state.get('app_lang', 'Français')
    placeholder = "Pose ta question sur l’entrepreneuriat" if lang == 'Français' else "Laaj sa laaj ci entrepreneuriat"
    user_msg = st.chat_input(placeholder)
    if user_msg:
        st.session_state['Fatouma_chat'].append({"role": "user", "content": user_msg})
        with st.chat_message("assistant"):
            response = Fatouma_chat_stream(st.session_state['Fatouma_chat'])
        st.session_state['Fatouma_chat'].append({"role": "assistant", "content": response})
        st.session_state['coaching_journal'].append({
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "question": user_msg,
            "reponse": response
        })

    # Rappel de statut sous le chat
    st.markdown("---")
    if st.session_state.get('profil_calcule'):
        st.success(tr('adja_profile_success'))
    else:
        st.info(tr('adja_info_prompt'))
        if st.button(tr('goto_eval_button'), key="goto_eval_button"):
            components.html("""
            <script>
            (function() {
              const doc = window.parent.document;
              const buttons = doc.querySelectorAll('button[role="tab"]');
              if (buttons && buttons.length > 0) {
                buttons[0].click();
                return;
              }
              const divs = doc.querySelectorAll('div[role="tab"]');
              if (divs && divs.length > 0) {
                divs[0].click();
              }
            })();
            </script>
            """, height=0)

    # Journal de coaching
    st.markdown("### " + tr('journal_coaching_title'))
    if st.session_state['coaching_journal']:
        df_journal = pd.DataFrame(st.session_state['coaching_journal'])
        st.dataframe(df_journal, use_container_width=True, hide_index=True)
        st.download_button(
            label=tr('download_journal_csv'),
            data=df_journal.to_csv(index=False),
            file_name=f"journal_coaching_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
            key="dl_journal_csv"
        )
    else:
        st.caption(tr('journal_empty_caption'))

st.markdown("---")
st.markdown(f"""
<div style='text-align: center'>
    <p>{tr('footer_tool_heading')}</p>
    <p style='font-size: 0.8em'>{tr('footer_tool_sub')}</p>
</div>
""", unsafe_allow_html=True)
